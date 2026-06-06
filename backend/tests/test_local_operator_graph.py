from pathlib import Path
import json
import os
import time

import pytest
from sqlmodel import select

from app.agent.graphs.memory_chat.nodes import _configured_local_operator_workspace_roots
from app.agent.graphs.memory_chat.nodes import _default_local_operator_workspace_roots
from app.agent.graphs.memory_chat.nodes import _build_react_agent_system_prompt
from app.agent.graphs.memory_chat.nodes import _infer_acceptance_criteria
from app.agent.graphs.local_operator.nodes import _build_local_operator_planner_prompt
from app.agent.graphs.local_operator.nodes import _rule_plan_action
from app.agent.graphs.local_operator.graph import get_local_operator_graph_mermaid
from app.agent.graphs.local_operator.graph import run_local_operator_graph
from app.core.config import settings
from app.local_operator.filesystem import LocalFilesystemService
from app.local_operator.command import LocalCommandExecutor, evaluate_command_policy
from app.local_operator.policy import LocalOperatorPolicy
from app.local_operator.remote import RemoteOperatorService
from app.local_operator.schemas import ToolResult
from app.local_operator.tools import create_read_tools
from app.models.agent_operation import AgentOperation


def test_read_file_blocks_path_escape(tmp_path: Path):
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    outside = tmp_path / "outside.txt"
    outside.write_text("secret", encoding="utf-8")
    service = LocalFilesystemService(LocalOperatorPolicy.from_roots([str(workspace)]))

    result = service.read_file(str(outside))

    assert result.ok is False
    assert result.error_code == "PATH_OUTSIDE_WORKSPACE"


def test_read_file_blocks_sensitive_file(tmp_path: Path):
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    (workspace / ".env").write_text("DASHSCOPE_API_KEY=secret", encoding="utf-8")
    service = LocalFilesystemService(LocalOperatorPolicy.from_roots([str(workspace)]))

    result = service.read_file(".env")

    assert result.ok is False
    assert result.error_code == "SENSITIVE_FILE_BLOCKED"
    assert result.blocked is True


def test_read_file_supports_line_range(tmp_path: Path):
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    (workspace / "note.txt").write_text("一\n二\n三\n四", encoding="utf-8")
    service = LocalFilesystemService(LocalOperatorPolicy.from_roots([str(workspace)]))

    result = service.read_file("note.txt", start_line=2, end_line=3)

    assert result.ok is True
    assert result.data["line_start"] == 2
    assert result.data["line_end"] == 3
    assert result.data["content"] == "二\n三"
    assert "     2\t二" in result.data["numbered_content"]


def test_read_file_strips_utf8_bom_and_normalizes_crlf(tmp_path: Path):
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    (workspace / "note.txt").write_bytes(b"\xef\xbb\xbfalpha\r\nbeta\r\n")
    service = LocalFilesystemService(LocalOperatorPolicy.from_roots([str(workspace)]))

    result = service.read_file("note.txt")

    assert result.ok is True
    assert result.data["content"] == "alpha\nbeta\n"
    assert result.data["total_bytes"] == 16
    assert result.data["modified_at"]


def test_read_file_supports_utf16le_bom(tmp_path: Path):
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    (workspace / "utf16.txt").write_bytes(b"\xff\xfe" + "你好\r\n世界".encode("utf-16-le"))
    service = LocalFilesystemService(LocalOperatorPolicy.from_roots([str(workspace)]))

    result = service.read_file("utf16.txt")

    assert result.ok is True
    assert result.data["content"] == "你好\n世界"


def test_read_file_rejects_null_byte_path(tmp_path: Path):
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    service = LocalFilesystemService(LocalOperatorPolicy.from_roots([str(workspace)]))

    result = service.read_file("bad\x00name.txt")

    assert result.ok is False
    assert result.error_code == "PATH_CONTAINS_NULL_BYTE"
    assert result.blocked is True


def test_read_file_blocks_dangerous_device_path(tmp_path: Path):
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    service = LocalFilesystemService(LocalOperatorPolicy.from_roots([str(workspace)]))

    result = service.read_file("/dev/zero")

    assert result.ok is False
    assert result.error_code == "DEVICE_PATH_BLOCKED"
    assert result.blocked is True


def test_read_document_extracts_docx_text(tmp_path: Path):
    from docx import Document

    workspace = tmp_path / "workspace"
    workspace.mkdir()
    document = Document()
    document.add_paragraph("项目背景")
    document.add_paragraph("AiMemo 需要读取 DOCX 文档。")
    document.save(workspace / "brief.docx")
    service = LocalFilesystemService(LocalOperatorPolicy.from_roots([str(workspace)]))

    result = service.read_document("brief.docx")

    assert result.ok is True
    assert result.data["document_type"] == "docx"
    assert result.data["relative_path"] == "brief.docx"
    assert "项目背景" in result.data["content"]
    assert "DOCX 文档" in result.data["content"]


def test_read_document_extracts_pdf_text(tmp_path: Path):
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    (workspace / "sample.pdf").write_bytes(
        b"%PDF-1.4\n"
        b"1 0 obj << /Type /Catalog /Pages 2 0 R >> endobj\n"
        b"2 0 obj << /Type /Pages /Kids [3 0 R] /Count 1 >> endobj\n"
        b"3 0 obj << /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >> endobj\n"
        b"4 0 obj << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> endobj\n"
        b"5 0 obj << /Length 44 >> stream\n"
        b"BT /F1 24 Tf 100 700 Td (Hello AiMemo PDF) Tj ET\n"
        b"endstream endobj\n"
        b"xref\n0 6\n0000000000 65535 f \n0000000009 00000 n \n"
        b"trailer << /Size 6 /Root 1 0 R >>\nstartxref\n405\n%%EOF\n"
    )
    service = LocalFilesystemService(LocalOperatorPolicy.from_roots([str(workspace)]))

    result = service.read_document("sample.pdf")

    assert result.ok is True
    assert result.data["document_type"] == "pdf"
    assert result.data["page_count"] == 1
    assert "Hello AiMemo PDF" in result.data["content"]


def test_read_document_rejects_legacy_doc(tmp_path: Path):
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    (workspace / "old.doc").write_bytes(b"not really a doc")
    service = LocalFilesystemService(LocalOperatorPolicy.from_roots([str(workspace)]))

    result = service.read_document("old.doc")

    assert result.ok is False
    assert result.error_code == "UNSUPPORTED_DOCUMENT_TYPE"
    assert result.blocked is True


def test_read_document_uses_same_sensitive_path_guard(tmp_path: Path):
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    (workspace / ".env").write_text("SECRET=1", encoding="utf-8")
    service = LocalFilesystemService(LocalOperatorPolicy.from_roots([str(workspace)]))

    result = service.read_document(".env")

    assert result.ok is False
    assert result.error_code == "SENSITIVE_FILE_BLOCKED"
    assert result.blocked is True


def test_read_document_tool_writes_audit(session, session_factory, tmp_path: Path):
    from docx import Document

    workspace = tmp_path / "workspace"
    workspace.mkdir()
    document = Document()
    document.add_paragraph("审计测试")
    document.save(workspace / "audit.docx")
    tools = create_read_tools(
        session_factory=session_factory,
        policy=LocalOperatorPolicy.from_roots([str(workspace)]),
        conversation_id=101,
        turn_id=202,
    )

    payload = json.loads(tools["read_document"].invoke({"path": "audit.docx"}))

    operations = session.exec(select(AgentOperation).where(AgentOperation.conversation_id == 101)).all()
    assert payload["ok"] is True
    assert "审计测试" in payload["data"]["content"]
    assert len(operations) == 1
    assert operations[0].tool_name == "read_document"
    assert operations[0].operation_type == "read"
    assert operations[0].status == "completed"


def test_rule_planner_uses_read_document_for_pdf_and_docx():
    pdf_action = _rule_plan_action("帮我读取 /tmp/report.pdf", workspace_roots=["."])
    docx_action = _rule_plan_action("查看文件 /tmp/brief.docx", workspace_roots=["."])

    assert pdf_action is not None
    assert pdf_action["tool_name"] == "read_document"
    assert docx_action is not None
    assert docx_action["tool_name"] == "read_document"


@pytest.mark.skipif(not str(Path.home()).startswith(str(Path.home().anchor)), reason="requires normal home path")
def test_policy_supports_home_expansion():
    home = Path.home().resolve()
    policy = LocalOperatorPolicy.from_roots([str(home)])

    resolved = policy.resolve_authorized_path("~")

    assert resolved == home


@pytest.mark.skipif(not (Path("C:/").exists()), reason="Windows drive path conversion only applies on Windows")
def test_policy_supports_posix_style_windows_drive_path():
    policy = LocalOperatorPolicy.from_roots(["C:/"])

    resolved = policy.resolve_authorized_path("/c/")

    assert resolved == Path("C:/").resolve()


@pytest.mark.skipif(not (Path("C:/").exists()), reason="Windows fixed drive path only applies on Windows")
def test_policy_allows_absolute_path_when_drive_root_is_authorized(tmp_path: Path):
    target = tmp_path / "readme.txt"
    target.write_text("hello from absolute path", encoding="utf-8")
    policy = LocalOperatorPolicy.from_roots([str(Path(target.anchor).resolve())])

    resolved = policy.resolve_authorized_path(str(target))

    assert resolved == target.resolve()


def test_read_file_large_file_allows_small_range(tmp_path: Path):
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    big_file = workspace / "large.txt"
    with big_file.open("w", encoding="utf-8", newline="\n") as file:
        for index in range(700_000):
            file.write(f"line-{index}\n")
    service = LocalFilesystemService(LocalOperatorPolicy.from_roots([str(workspace)]))

    result = service.read_file("large.txt", start_line=700_000, end_line=700_000, max_bytes=128)

    assert result.ok is True
    assert result.data["line_start"] == 700_000
    assert result.data["line_end"] == 700_000
    assert result.data["content"] == "line-699999"


def test_missing_path_returns_nearby_suggestion(tmp_path: Path):
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    (workspace / "memory-chat-graph.md").write_text("graph", encoding="utf-8")
    service = LocalFilesystemService(LocalOperatorPolicy.from_roots([str(workspace)]))

    result = service.read_file("memory-chat.md")

    assert result.ok is False
    assert result.error_code == "PATH_NOT_FOUND"
    assert "memory-chat-graph.md" in result.message


def test_search_text_finds_matches(tmp_path: Path):
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    (workspace / "a.py").write_text("alpha\nmemory_key = 'x'\nomega", encoding="utf-8")
    service = LocalFilesystemService(LocalOperatorPolicy.from_roots([str(workspace)]))

    result = service.search_text(".", query="memory_key", include_glob="*.py")

    assert result.ok is True
    assert result.data["match_count"] == 1
    assert result.data["matches"][0]["relative_path"] == "a.py"
    assert result.data["matches"][0]["line"] == 2


def test_write_file_creates_parent_directories(tmp_path: Path):
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    service = LocalFilesystemService(LocalOperatorPolicy.from_roots([str(workspace)]))

    result = service.write_file("docs/hello.txt", content="你好，Memo。")

    assert result.ok is True
    assert result.data["type"] == "create"
    assert result.data["relative_path"] == "docs/hello.txt"
    assert (workspace / "docs" / "hello.txt").read_text(encoding="utf-8") == "你好，Memo。"


def test_write_file_requires_read_before_overwrite(tmp_path: Path):
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    (workspace / "hello.txt").write_text("old", encoding="utf-8")
    service = LocalFilesystemService(LocalOperatorPolicy.from_roots([str(workspace)]))

    blocked = service.write_file("hello.txt", content="new", overwrite=True)
    read_result = service.read_file("hello.txt")
    allowed = service.write_file("hello.txt", content="new", overwrite=True)

    assert blocked.ok is False
    assert blocked.error_code == "READ_BEFORE_WRITE_REQUIRED"
    assert blocked.blocked is True
    assert read_result.ok is True
    assert allowed.ok is True
    assert allowed.data["type"] == "update"
    assert (workspace / "hello.txt").read_text(encoding="utf-8") == "new"


def test_write_file_requires_full_read_before_overwrite(tmp_path: Path):
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    target = workspace / "hello.txt"
    target.write_text("old\ncontent", encoding="utf-8")
    service = LocalFilesystemService(LocalOperatorPolicy.from_roots([str(workspace)]))

    partial_read = service.read_file("hello.txt", start_line=1, end_line=1)
    result = service.write_file("hello.txt", content="new", overwrite=True)

    assert partial_read.ok is True
    assert partial_read.data["full_view"] is False
    assert result.ok is False
    assert result.error_code == "WRITE_WITH_PARTIAL_READ"
    assert target.read_text(encoding="utf-8") == "old\ncontent"


def test_write_file_allows_confirmed_overwrite_after_partial_read(tmp_path: Path):
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    target = workspace / "hello.txt"
    target.write_text("old\ncontent", encoding="utf-8")
    service = LocalFilesystemService(LocalOperatorPolicy.from_roots([str(workspace)]))

    partial_read = service.read_file("hello.txt", start_line=1, end_line=1)
    result = service.write_file(
        "hello.txt",
        content="new",
        overwrite=True,
        confirmed_overwrite_without_read=True,
    )

    assert partial_read.ok is True
    assert partial_read.data["full_view"] is False
    assert result.ok is True
    assert result.data["type"] == "update"
    assert result.data["bypassed_read_before_write"] is True
    assert target.read_text(encoding="utf-8") == "new"


def test_write_file_keeps_sensitive_block_with_confirmed_overwrite(tmp_path: Path):
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    target = workspace / ".env"
    target.write_text("SECRET=old", encoding="utf-8")
    service = LocalFilesystemService(LocalOperatorPolicy.from_roots([str(workspace)]))

    result = service.write_file(
        ".env",
        content="SECRET=new",
        overwrite=True,
        confirmed_overwrite_without_read=True,
    )

    assert result.ok is False
    assert result.error_code == "SENSITIVE_FILE_BLOCKED"
    assert result.blocked is True
    assert target.read_text(encoding="utf-8") == "SECRET=old"


def test_write_file_rejects_changed_file_after_read(tmp_path: Path):
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    target = workspace / "hello.txt"
    target.write_text("old", encoding="utf-8")
    service = LocalFilesystemService(LocalOperatorPolicy.from_roots([str(workspace)]))

    read_result = service.read_file("hello.txt")
    target.write_text("changed outside", encoding="utf-8")
    future = time.time() + 5
    os.utime(target, (future, future))
    result = service.write_file("hello.txt", content="new", overwrite=True)

    assert read_result.ok is True
    assert read_result.data["full_view"] is True
    assert result.ok is False
    assert result.error_code == "FILE_MTIME_CHANGED"
    assert target.read_text(encoding="utf-8") == "changed outside"


def test_write_file_rejects_placeholder_content(tmp_path: Path):
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    service = LocalFilesystemService(LocalOperatorPolicy.from_roots([str(workspace)]))

    result = service.write_file("review.md", content="## 对用户的评价\n\n此处填写对用户的评价内容。")

    assert result.ok is False
    assert result.error_code == "PLACEHOLDER_CONTENT_REJECTED"
    assert result.blocked is True
    assert not (workspace / "review.md").exists()


def test_exec_command_runs_short_readonly_command(tmp_path: Path):
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    executor = LocalCommandExecutor(LocalOperatorPolicy.from_roots([str(workspace)]))

    result = executor.exec_command(command="python --version", cwd=".")

    assert result.ok is True
    assert result.tool_name == "exec_command"
    assert result.data["exit_code"] == 0
    assert "Python" in (result.data["stdout"] + result.data["stderr"])


def test_exec_command_timeout_defaults_come_from_project_config():
    assert settings.local_operator_exec_default_timeout_ms >= 180_000
    assert settings.local_operator_exec_max_timeout_ms >= settings.local_operator_exec_default_timeout_ms


def test_exec_command_reports_non_zero_exit_as_failure(tmp_path: Path):
    """非 0 退出码必须反馈为失败，供 agent 根据 stderr 触发修复或重规划。"""

    workspace = tmp_path / "workspace"
    workspace.mkdir()
    executor = LocalCommandExecutor(LocalOperatorPolicy.from_roots([str(workspace)]))

    result = executor.exec_command(command="python -c \"import sys; sys.exit(7)\"", cwd=".")

    assert result.ok is False
    assert result.blocked is False
    assert result.error_code == "COMMAND_EXITED_NON_ZERO"
    assert result.data["exit_code"] == 7


def test_exec_command_turns_pipe_broken_into_tool_failure(tmp_path: Path, monkeypatch):
    """Windows 管道断开不应把整个 agent loop 炸掉。"""

    workspace = tmp_path / "workspace"
    workspace.mkdir()
    executor = LocalCommandExecutor(LocalOperatorPolicy.from_roots([str(workspace)]))

    class FakeProc:
        returncode = None

        def communicate(self, timeout=None):  # noqa: ARG002
            raise OSError(233, "管道的另一端上无任何进程。")

        def poll(self):
            return 0

    monkeypatch.setattr("app.local_operator.command._spawn_subprocess", lambda *args, **kwargs: FakeProc())

    result = executor.exec_command(command="python --version", cwd=".")

    assert result.ok is False
    assert result.error_code == "COMMAND_PIPE_BROKEN"
    assert result.blocked is False
    assert result.data["pipe_broken"] is True
    assert "管道断开" in result.message


def test_exec_command_blocks_destructive_command(tmp_path: Path):
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    executor = LocalCommandExecutor(LocalOperatorPolicy.from_roots([str(workspace)]))

    result = executor.exec_command(command="git reset --hard", cwd=".")

    assert result.ok is False
    assert result.error_code == "COMMAND_BLOCKED"
    assert result.blocked is True


def test_exec_command_blocks_raw_scp_and_points_to_remote_tools():
    decision = evaluate_command_policy("scp index.html root@example.com:/usr/share/nginx/html/index.html")

    assert decision.allowed is False
    assert decision.risk_level == "high"
    assert "远程操作工具" in decision.reason


def test_create_read_tools_exposes_remote_tools(session_factory, tmp_path: Path):
    workspace = tmp_path / "workspace"
    workspace.mkdir()

    tools = create_read_tools(
        session_factory=session_factory,
        policy=LocalOperatorPolicy.from_roots([str(workspace)]),
        conversation_id=101,
        turn_id=202,
    )

    assert "remote_connectivity_check" in tools
    assert "remote_upload_file" in tools
    assert "remote_exec" in tools
    assert "remote_verify_http" in tools


def test_remote_connectivity_classifies_interactive_auth(tmp_path: Path, monkeypatch):
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    service = RemoteOperatorService(LocalOperatorPolicy.from_roots([str(workspace)]))

    monkeypatch.setattr("app.local_operator.remote.shutil.which", lambda name: name)
    monkeypatch.setattr(
        "app.local_operator.remote._run_process",
        lambda *args, **kwargs: ToolResult(
            ok=False,
            tool_name="remote_process",
            data={
                "stdout": "",
                "stderr": "Permission denied (publickey,password).",
                "timed_out": False,
                "exit_code": 255,
            },
            error_code="REMOTE_PROCESS_EXITED_NON_ZERO",
            message="远程进程命令以非 0 状态退出。",
        ),
    )

    result = service.connectivity_check(host="example.com", username="root")

    assert result.ok is False
    assert result.tool_name == "remote_connectivity_check"
    assert result.error_code == "INTERACTIVE_AUTH_REQUIRED"
    assert result.blocked is True


def test_memory_chat_prompt_requires_remote_tools_for_server_operations():
    prompt = _build_react_agent_system_prompt()

    assert "remote_connectivity_check" in prompt
    assert "remote_upload_file" in prompt
    assert "不要把 ssh/scp" in prompt


def test_remote_acceptance_criteria_requires_tool_and_verification():
    criteria = _infer_acceptance_criteria("登录远程服务器，把静态页面上传到 nginx 并部署")
    text = "\n".join(criteria)

    assert "remote_* 工具" in text
    assert "remote_upload_file/remote_exec" in text
    assert "remote_verify_http" in text


def test_local_operator_graph_reads_file_and_writes_audit(
    session,
    session_factory,
    tmp_path: Path,
):
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    (workspace / "hello.txt").write_text("你好，Ai 记。", encoding="utf-8")

    result = run_local_operator_graph(
        session_factory=session_factory,
        conversation_id=7,
        turn_id=11,
        user_input="帮我读取 hello.txt 这个文件",
        workspace_roots=[str(workspace)],
    )

    operations = session.exec(select(AgentOperation)).all()
    assert result["needs_tool"] is True
    assert "你好，Ai 记。" in result["final_answer"]
    assert len(operations) == 1
    assert operations[0].conversation_id == 7
    assert operations[0].turn_id == 11
    assert operations[0].operation_type == "read"
    assert operations[0].status == "completed"
    assert operations[0].tool_name == "read_file"


def test_local_operator_graph_creates_file_and_writes_audit(
    session,
    session_factory,
    tmp_path: Path,
):
    workspace = tmp_path / "workspace"
    workspace.mkdir()

    result = run_local_operator_graph(
        session_factory=session_factory,
        conversation_id=19,
        turn_id=23,
        user_input="请创建文件 hello.txt 内容是 你好，AiMemo。",
        workspace_roots=[str(workspace)],
    )

    operations = session.exec(select(AgentOperation).where(AgentOperation.conversation_id == 19)).all()
    assert result["needs_tool"] is True
    assert "已创建 `hello.txt`" in result["final_answer"]
    assert (workspace / "hello.txt").read_text(encoding="utf-8") == "你好，AiMemo"
    assert len(operations) == 1
    assert operations[0].operation_type == "write"
    assert operations[0].status == "completed"
    assert operations[0].tool_name == "write_file"
    assert operations[0].risk_level == "medium"


def test_local_operator_graph_executes_command_and_writes_audit(
    session,
    session_factory,
    tmp_path: Path,
):
    workspace = tmp_path / "workspace"
    workspace.mkdir()

    result = run_local_operator_graph(
        session_factory=session_factory,
        conversation_id=31,
        turn_id=41,
        user_input="帮我执行命令 `python --version`",
        workspace_roots=[str(workspace)],
    )

    operations = session.exec(select(AgentOperation).where(AgentOperation.conversation_id == 31)).all()
    assert result["needs_tool"] is True
    assert "python --version" in result["final_answer"]
    assert len(operations) == 1
    assert operations[0].operation_type == "exec"
    assert operations[0].tool_name == "exec_command"
    assert operations[0].status == "completed"


def test_local_operator_graph_does_not_create_placeholder_review_file(
    session,
    session_factory,
    tmp_path: Path,
):
    workspace = tmp_path / "workspace"
    workspace.mkdir()

    result = run_local_operator_graph(
        session_factory=session_factory,
        conversation_id=21,
        turn_id=25,
        user_input="请在 review.md 创建一个 markdown 文件用于写对我的评价，内容是 此处填写对用户的评价内容。",
        workspace_roots=[str(workspace)],
    )

    operations = session.exec(select(AgentOperation).where(AgentOperation.conversation_id == 21)).all()
    assert result["needs_tool"] is True
    assert "PLACEHOLDER_CONTENT_REJECTED" in result["final_answer"]
    assert not (workspace / "review.md").exists()
    assert len(operations) == 1
    assert operations[0].status == "blocked"
    assert operations[0].tool_name == "write_file"


def test_local_operator_graph_overwrites_existing_file_after_info_check(
    session,
    session_factory,
    tmp_path: Path,
):
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    (workspace / "hello.txt").write_text("old", encoding="utf-8")

    result = run_local_operator_graph(
        session_factory=session_factory,
        conversation_id=20,
        turn_id=24,
        user_input="请覆盖写入 hello.txt 内容是 new-content",
        workspace_roots=[str(workspace)],
    )

    operations = session.exec(select(AgentOperation).where(AgentOperation.conversation_id == 20).order_by(AgentOperation.id)).all()
    assert result["needs_tool"] is True
    assert (workspace / "hello.txt").read_text(encoding="utf-8") == "new-content"
    assert [operation.tool_name for operation in operations] == ["read_file", "write_file"]
    assert [operation.operation_type for operation in operations] == ["read", "write"]


def test_local_operator_graph_detects_project_existence_question(
    session,
    session_factory,
    tmp_path: Path,
):
    workspace = tmp_path / "Ai记"
    workspace.mkdir()

    result = run_local_operator_graph(
        session_factory=session_factory,
        conversation_id=8,
        turn_id=12,
        user_input="我当前电脑有没有 Ai记 的项目？",
        workspace_roots=[str(workspace)],
    )

    operations = session.exec(select(AgentOperation).where(AgentOperation.conversation_id == 8)).all()
    assert result["needs_tool"] is True
    assert "Ai记" in result["final_answer"]
    assert len(operations) == 1
    assert operations[0].tool_name == "search_files"


def test_local_operator_graph_searches_home_for_whole_computer_project_question(
    session,
    session_factory,
    tmp_path: Path,
):
    home = tmp_path / "home"
    project = home / "AiMemo"
    project.mkdir(parents=True)

    result = run_local_operator_graph(
        session_factory=session_factory,
        conversation_id=18,
        turn_id=22,
        user_input="我当前电脑有没有 AiMemo 这个项目？",
        workspace_roots=[str(tmp_path / "repo"), str(home)],
    )

    operations = session.exec(select(AgentOperation).where(AgentOperation.conversation_id == 18)).all()
    assert result["needs_tool"] is True
    assert "AiMemo" in result["final_answer"]
    assert len(operations) == 1
    assert operations[0].tool_name == "search_files"


def test_local_operator_graph_uses_llm_planner_for_ambiguous_local_request(
    session,
    session_factory,
    tmp_path: Path,
):
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    planner_calls: list[str] = []

    def fake_planner(user_input: str):
        planner_calls.append(user_input)
        return {
            "tool_name": "get_file_info",
            "arguments": {"path": "."},
            "reason": "测试 planner 判断为本地项目确认。",
        }

    result = run_local_operator_graph(
        session_factory=session_factory,
        conversation_id=9,
        turn_id=13,
        user_input="帮我确认这个仓库还在工作区里吗？",
        workspace_roots=[str(workspace)],
        planner=fake_planner,
    )

    operations = session.exec(select(AgentOperation).where(AgentOperation.conversation_id == 9)).all()
    assert planner_calls == ["帮我确认这个仓库还在工作区里吗？"]
    assert result["needs_tool"] is True
    assert operations[0].tool_name == "get_file_info"


def test_local_operator_graph_skips_planner_for_plain_chat(
    session_factory,
    tmp_path: Path,
):
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    planner_calls: list[str] = []

    def fake_planner(user_input: str):
        planner_calls.append(user_input)
        return {
            "tool_name": "get_file_info",
            "arguments": {"path": "."},
            "reason": "不应该被调用。",
        }

    result = run_local_operator_graph(
        session_factory=session_factory,
        conversation_id=10,
        turn_id=14,
        user_input="你觉得我今天状态怎么样？",
        workspace_roots=[str(workspace)],
        planner=fake_planner,
    )

    assert planner_calls == []
    assert result["needs_tool"] is False


def test_local_operator_graph_mermaid_shows_read_write_tool_routes(session_factory):
    mermaid = get_local_operator_graph_mermaid(session_factory=session_factory)

    assert "plan_tool_use" in mermaid
    assert "select_tool" in mermaid
    assert "run_read_tool" in mermaid
    assert "run_write_tool" in mermaid
    assert "run_exec_tool" in mermaid
    assert "plan_read_operation" not in mermaid
    assert "select_read_tool" not in mermaid


def test_default_local_operator_workspace_roots_include_repo_and_home():
    roots = [Path(root) for root in _default_local_operator_workspace_roots()]

    assert Path.home().resolve() in roots
    assert any((root / "backend").exists() and (root / "frontend").exists() for root in roots)


@pytest.mark.skipif(not (Path("C:/").exists()), reason="Windows fixed drive path only applies on Windows")
def test_default_local_operator_workspace_roots_include_fixed_drive_roots():
    roots = [Path(root).resolve() for root in _default_local_operator_workspace_roots()]

    assert Path("C:/").resolve() in roots


def test_local_operator_planner_prompt_does_not_preemptively_refuse_absolute_paths():
    prompt = _build_local_operator_planner_prompt("帮我看看 C:\\Windows\\Logs\\setup.log")

    assert "用户明确给出的绝对路径可以原样传给工具" in prompt
    assert "不能凭空声称" in prompt


def test_configured_local_operator_workspace_roots_parse_semicolon_and_comma(monkeypatch):
    monkeypatch.setattr(settings, "local_operator_workspace_roots", "D:\\资料;~/Documents,E:\\Projects")

    roots = _configured_local_operator_workspace_roots()

    assert roots == ["D:\\资料", "~/Documents", "E:\\Projects"]
