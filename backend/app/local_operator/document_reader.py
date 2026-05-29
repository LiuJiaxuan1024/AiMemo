from __future__ import annotations

from pathlib import Path
from typing import Any

from app.local_operator.policy import LocalOperatorPolicy
from app.local_operator.schemas import ToolResult


SUPPORTED_DOCUMENT_EXTENSIONS = {".pdf", ".docx"}
MAX_DOCUMENT_BYTES = 25 * 1024 * 1024


def read_document(
    path: Path,
    policy: LocalOperatorPolicy,
    *,
    max_chars: int = 80_000,
) -> ToolResult:
    """Extract readable text from supported document files.

    This intentionally lives beside, not inside, read_file. read_file remains a
    text/source-code reader with read-before-write semantics; read_document is
    a lossy, read-only extraction surface for office-style documents.
    """

    if policy.is_sensitive_path(path):
        return _error("SENSITIVE_FILE_BLOCKED", "该文件可能包含敏感信息，已拒绝读取。", blocked=True)
    if path.is_dir():
        return _error("PATH_IS_DIRECTORY", "路径是目录，不是文件。")

    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_DOCUMENT_EXTENSIONS:
        if suffix == ".doc":
            return _error(
                "UNSUPPORTED_DOCUMENT_TYPE",
                "暂不直接解析旧版 .doc 文件。请先用 LibreOffice/Word 转成 .docx 或 PDF 后再读取。",
                blocked=True,
            )
        return _error(
            "UNSUPPORTED_DOCUMENT_TYPE",
            "暂只支持 PDF 和 DOCX 文档解析；普通文本/源码请使用 read_file。",
            blocked=True,
        )

    stat = path.stat()
    if stat.st_size > MAX_DOCUMENT_BYTES:
        return _error(
            "DOCUMENT_TOO_LARGE",
            f"文档过大（{stat.st_size} bytes），当前 read_document 上限为 {MAX_DOCUMENT_BYTES} bytes。",
            blocked=True,
        )

    max_chars = min(max(int(max_chars or 80_000), 1), 200_000)
    if suffix == ".pdf":
        return _read_pdf(path, policy, stat.st_size, max_chars=max_chars)
    return _read_docx(path, policy, stat.st_size, max_chars=max_chars)


def _read_pdf(path: Path, policy: LocalOperatorPolicy, total_bytes: int, *, max_chars: int) -> ToolResult:
    try:
        from pypdf import PdfReader
        from pypdf.errors import PdfReadError
    except ImportError:
        return _error(
            "DOCUMENT_READER_DEPENDENCY_MISSING",
            "缺少 pypdf 依赖，无法解析 PDF。请先安装后端依赖。",
        )

    try:
        reader = PdfReader(str(path))
        if reader.is_encrypted:
            try:
                reader.decrypt("")
            except Exception:
                return _error("PASSWORD_PROTECTED_DOCUMENT", "PDF 已加密，请提供未加密版本。", blocked=True)
        pages = []
        truncated = False
        for page_index, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            if text.strip():
                pages.append(f"## Page {page_index}\n\n{text.strip()}")
            if sum(len(item) for item in pages) >= max_chars:
                truncated = True
                break
        content = "\n\n".join(pages).strip()
        if len(content) > max_chars:
            content = content[:max_chars]
            truncated = True
        if not content:
            return _error(
                "DOCUMENT_TEXT_EMPTY",
                "未能从 PDF 中提取到文本。它可能是扫描件或图片型 PDF，后续需要 OCR 能力。",
            )
        return _ok(
            {
                "path": path.as_posix(),
                "relative_path": policy.relative_path(path),
                "document_type": "pdf",
                "total_bytes": total_bytes,
                "page_count": len(reader.pages),
                "char_count": len(content),
                "truncated": truncated,
                "extraction_method": "pypdf",
                "content": content,
            }
        )
    except PdfReadError as exc:
        return _error("DOCUMENT_PARSE_FAILED", f"PDF 解析失败：{exc}")
    except Exception as exc:
        return _error("DOCUMENT_PARSE_FAILED", f"PDF 解析失败：{exc}")


def _read_docx(path: Path, policy: LocalOperatorPolicy, total_bytes: int, *, max_chars: int) -> ToolResult:
    try:
        from docx import Document
    except ImportError:
        return _error(
            "DOCUMENT_READER_DEPENDENCY_MISSING",
            "缺少 python-docx 依赖，无法解析 DOCX。请先安装后端依赖。",
        )

    try:
        document = Document(str(path))
        parts: list[str] = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                parts.append(text)
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        content = "\n\n".join(parts).strip()
        truncated = False
        if len(content) > max_chars:
            content = content[:max_chars]
            truncated = True
        if not content:
            return _error("DOCUMENT_TEXT_EMPTY", "未能从 DOCX 中提取到文本。")
        return _ok(
            {
                "path": path.as_posix(),
                "relative_path": policy.relative_path(path),
                "document_type": "docx",
                "total_bytes": total_bytes,
                "page_count": None,
                "char_count": len(content),
                "truncated": truncated,
                "extraction_method": "python-docx",
                "content": content,
            }
        )
    except Exception as exc:
        return _error("DOCUMENT_PARSE_FAILED", f"DOCX 解析失败：{exc}")


def _ok(data: dict[str, Any]) -> ToolResult:
    return ToolResult(ok=True, tool_name="read_document", data=data)


def _error(error_code: str, message: str, *, blocked: bool = False) -> ToolResult:
    return ToolResult(
        ok=False,
        tool_name="read_document",
        error_code=error_code,
        message=message,
        blocked=blocked,
    )
