from __future__ import annotations

from pathlib import Path

from app.rag.document_parsers.base import DocumentBlock, DocumentImageAsset, DocumentParseError, ParsedDocument


def parse_docx(path: Path) -> ParsedDocument:
    try:
        from docx import Document
    except ImportError as exc:
        raise DocumentParseError("DOCUMENT_READER_DEPENDENCY_MISSING", "缺少 python-docx 依赖，无法解析 DOCX。") from exc

    try:
        document = Document(str(path))
    except Exception as exc:
        raise DocumentParseError("DOCUMENT_PARSE_FAILED", f"DOCX 解析失败：{exc}") from exc

    blocks: list[DocumentBlock] = []
    image_assets: list[DocumentImageAsset] = []
    heading_stack: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = (paragraph.style.name if paragraph.style is not None else "").lower()
        if style_name.startswith("heading"):
            level = _heading_level(style_name)
            heading_stack = heading_stack[: level - 1]
            heading_stack.append(text)
            blocks.append(
                DocumentBlock(
                    text=text,
                    block_type="heading",
                    heading_path=list(heading_stack),
                    metadata={"level": level},
                )
            )
        else:
            blocks.append(DocumentBlock(text=text, block_type="paragraph", heading_path=list(heading_stack)))

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                blocks.append(DocumentBlock(text=" | ".join(cells), block_type="table", heading_path=list(heading_stack)))

    for image_index, shape in enumerate(document.inline_shapes, start=1):
        width = getattr(shape, "width", None)
        height = getattr(shape, "height", None)
        image_assets.append(
            DocumentImageAsset(
                asset_id=f"docx-image-{image_index}",
                parser="docx",
                location_label=f"DOCX 图片 {image_index}",
                heading_path=list(heading_stack),
                source_offset=image_index,
                data=_inline_shape_blob(document, shape),
                mime_type=_inline_shape_mime_type(document, shape),
                width=int(width) if width is not None else None,
                height=int(height) if height is not None else None,
            )
        )

    if not blocks:
        if not image_assets:
            raise DocumentParseError("DOCUMENT_TEXT_EMPTY", "未能从 DOCX 中提取到文本。")
    title = next((block.text for block in blocks if block.block_type == "heading"), None)
    return ParsedDocument(parser="docx", blocks=blocks, title=title or path.stem, image_assets=image_assets)


def _heading_level(style_name: str) -> int:
    digits = "".join(char for char in style_name if char.isdigit())
    if not digits:
        return 1
    return max(1, min(int(digits), 6))


def _inline_shape_part(document, shape):
    try:
        blip = shape._inline.graphic.graphicData.pic.blipFill.blip
        rel_id = blip.embed
        return document.part.related_parts.get(rel_id)
    except Exception:
        return None


def _inline_shape_blob(document, shape) -> bytes | None:
    part = _inline_shape_part(document, shape)
    return getattr(part, "blob", None) if part is not None else None


def _inline_shape_mime_type(document, shape) -> str | None:
    part = _inline_shape_part(document, shape)
    return getattr(part, "content_type", None) if part is not None else None
